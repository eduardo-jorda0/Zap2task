"""Exportação da timeline final em `.txt`, `.json` e `.docx` (RF6).

Todos os exportadores recebem o mesmo `RelatorioFinal` (timeline + metadados +
relatórios de duplicatas/falhas), garantindo que as três saídas sempre tragam
a mesma informação, só que formatada de forma diferente.
"""

import json
from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument

import config
from models.tipos import RelatorioFinal, TimelineItem, TipoItemTimeline
from utils.logger import get_logger

logger = get_logger(__name__)

_ROTULOS_TIPO = {
    TipoItemTimeline.TEXTO: "texto",
    TipoItemTimeline.AUDIO: "áudio",
    TipoItemTimeline.DOCUMENTO: "documento",
}
_AVISO_BAIXA_CONFIANCA = "[transcrição de baixa confiança — revisar manualmente] "
_LARGURA_SEPARADOR = 56


def exportar_txt(relatorio: RelatorioFinal, destino_dir: Path) -> Path:
    """Gera um arquivo `.txt` legível com a conversa completa mesclada e ordenada (RF6.1).

    Args:
        relatorio: pacote completo de dados da execução.
        destino_dir: diretório onde o arquivo será salvo.

    Returns:
        Caminho do arquivo `.txt` gerado.
    """
    caminho_saida = destino_dir / _nome_arquivo("conversa", "txt", relatorio.metadados.processado_em)
    caminho_saida.write_text(_montar_conteudo_txt(relatorio), encoding="utf-8")
    logger.info("Exportado .txt: %s", caminho_saida)
    return caminho_saida


def exportar_json(relatorio: RelatorioFinal, destino_dir: Path) -> Path:
    """Gera um arquivo `.json` estruturado com timeline, metadados e relatórios (RF6.2).

    Args:
        relatorio: pacote completo de dados da execução.
        destino_dir: diretório onde o arquivo será salvo.

    Returns:
        Caminho do arquivo `.json` gerado.
    """
    caminho_saida = destino_dir / _nome_arquivo("conversa", "json", relatorio.metadados.processado_em)
    dados = _montar_dados_json(relatorio)
    caminho_saida.write_text(json.dumps(dados, ensure_ascii=False, indent=2), encoding="utf-8")
    logger.info("Exportado .json: %s", caminho_saida)
    return caminho_saida


def exportar_docx(relatorio: RelatorioFinal, destino_dir: Path) -> Path:
    """Gera um documento Word formatado, para revisão manual (RF6.3, opcional).

    Args:
        relatorio: pacote completo de dados da execução.
        destino_dir: diretório onde o arquivo será salvo.

    Returns:
        Caminho do arquivo `.docx` gerado.
    """
    caminho_saida = destino_dir / _nome_arquivo("conversa", "docx", relatorio.metadados.processado_em)
    documento = _montar_documento_docx(relatorio)
    documento.save(str(caminho_saida))
    logger.info("Exportado .docx: %s", caminho_saida)
    return caminho_saida


def _nome_arquivo(prefixo: str, extensao: str, processado_em: datetime) -> str:
    """Nomeia o arquivo de saída com o timestamp da execução, evitando sobrescrever resultados anteriores (RF6.4)."""
    return f"{prefixo}_{processado_em.strftime('%Y%m%d_%H%M%S')}.{extensao}"


def _montar_conteudo_txt(relatorio: RelatorioFinal) -> str:
    metadados = relatorio.metadados
    linhas = [
        "=== Zap2Task Audio Engine — Relatório de Conversa ===",
        f"Processado em: {metadados.processado_em.strftime('%d/%m/%Y %H:%M')}",
        f"Total de mensagens de texto: {metadados.total_mensagens_texto}",
        f"Total de áudios transcritos: {metadados.total_audios_transcritos}",
        f"Duplicatas removidas: {metadados.duplicatas_removidas}",
        "=" * _LARGURA_SEPARADOR,
    ]
    linhas.extend(_formatar_linha_txt(item) for item in relatorio.timeline)
    return "\n".join(linhas) + "\n"


def _formatar_linha_txt(item: TimelineItem) -> str:
    cabecalho = _formatar_cabecalho(item)
    conteudo = _prefixo_baixa_confianca(item) + item.conteudo

    primeira_linha, *continuacoes = conteudo.split("\n")
    linhas = [cabecalho + primeira_linha]
    linhas.extend(f"    {linha}" for linha in continuacoes)
    return "\n".join(linhas)


def _formatar_cabecalho(item: TimelineItem) -> str:
    rotulo_tipo = _ROTULOS_TIPO[item.tipo]
    sufixo_confianca = ""
    if item.tipo == TipoItemTimeline.AUDIO and item.confianca is not None:
        sufixo_confianca = f", confiança {item.confianca:.0%}"
    timestamp = item.timestamp.strftime("%d/%m/%Y %H:%M")
    return f"[{timestamp}] {item.remetente} ({rotulo_tipo}{sufixo_confianca}): "


def _prefixo_baixa_confianca(item: TimelineItem) -> str:
    """Marca transcrições de baixa confiança para revisão manual (RF4.5)."""
    if (
        item.tipo == TipoItemTimeline.AUDIO
        and item.confianca is not None
        and item.confianca < config.LIMIAR_CONFIANCA_BAIXA
    ):
        return _AVISO_BAIXA_CONFIANCA
    return ""


def _montar_dados_json(relatorio: RelatorioFinal) -> dict:
    metadados = relatorio.metadados
    return {
        "metadados": {
            "processado_em": metadados.processado_em.isoformat(),
            "total_mensagens_texto": metadados.total_mensagens_texto,
            "total_audios_transcritos": metadados.total_audios_transcritos,
            "duplicatas_removidas": metadados.duplicatas_removidas,
            "modelo_whisper_usado": metadados.modelo_whisper_usado,
        },
        "timeline": [_item_para_dict(item) for item in relatorio.timeline],
        "duplicatas_removidas": [
            {"arquivo": duplicata.arquivo_removido, "duplicata_de": duplicata.duplicata_de}
            for duplicata in relatorio.duplicatas_removidas
        ],
        "falhas_conversao": [
            {"arquivo": falha.arquivo, "motivo": falha.motivo} for falha in relatorio.falhas_conversao
        ],
        "falhas_transcricao": [
            {"arquivo": falha.arquivo, "motivo": falha.motivo} for falha in relatorio.falhas_transcricao
        ],
    }


def _item_para_dict(item: TimelineItem) -> dict:
    return {
        "timestamp": item.timestamp.isoformat(),
        "remetente": item.remetente,
        "tipo": item.tipo.value,
        "conteudo": item.conteudo,
        "confianca": item.confianca,
        "arquivo_origem": item.arquivo_origem,
    }


def _montar_documento_docx(relatorio: RelatorioFinal) -> DocxDocument:
    metadados = relatorio.metadados
    documento = DocxDocument()

    documento.add_heading("Zap2Task Audio Engine — Relatório de Conversa", level=1)
    documento.add_paragraph(f"Processado em: {metadados.processado_em.strftime('%d/%m/%Y %H:%M')}")
    documento.add_paragraph(f"Total de mensagens de texto: {metadados.total_mensagens_texto}")
    documento.add_paragraph(f"Total de áudios transcritos: {metadados.total_audios_transcritos}")
    documento.add_paragraph(f"Duplicatas removidas: {metadados.duplicatas_removidas}")

    for item in relatorio.timeline:
        _adicionar_paragrafo_item(documento, item)

    return documento


def _adicionar_paragrafo_item(documento: DocxDocument, item: TimelineItem) -> None:
    paragrafo = documento.add_paragraph()
    execucao_cabecalho = paragrafo.add_run(_formatar_cabecalho(item))
    execucao_cabecalho.bold = True
    paragrafo.add_run(_prefixo_baixa_confianca(item) + item.conteudo)
